"""
Mission Logic & DOP/DIP Parser
----------------------------
A comprehensive tool for mission-critical data extraction from technical documents.
Parses complex logical expressions and DOP/DIP (Discrete Output/Input Point) data.

Key features:
- Parses nested logical expressions with AND/OR operators
- Normalizes unicode characters, units, and bracketing
- Identifies descriptive blocks vs. mathematical expressions
- Extracts DOP/DIP data with subsystem and bus mapping
- Handles both Word (.docx) and PDF documents
- Outputs to CSV, Excel, and JSON formats
"""

import re
import pandas as pd
import json
import fitz
from docx import Document
from docx.shared import Pt
from pdf2docx import Converter
import tempfile
import os
from IPython.display import display
import openpyxl
from tkinter import Tk, filedialog

# -----------------------
# Character Normalization
# ------------------------

# Mapping dictionaries for special characters
SUPERSCRIPT_MAP = str.maketrans("⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻", "0123456789+-")
SUBSCRIPT_MAP = str.maketrans("₀₁₂₃₄₅₆₇₈₉", "0123456789")
UNICODE_SUBSCRIPT_MAP = {
    # Subscripts
    'ₐ': 'a', 'ₑ': 'e', 'ₒ': 'o', 'ₓ': 'x', 'ₕ': 'h',
    'ₖ': 'k', 'ₗ': 'l', 'ₘ': 'm', 'ₙ': 'n', 'ₚ': 'p',
    'ₛ': 's', 'ₜ': 't', 'ᵣ': 'r', 'ᵤ': 'u', 'ᵥ': 'v',
    'ₓ': 'x', 'ᵧ': 'y', '₂': '2', '₃': '3', '₄': '4',
    '₅': '5', '₆': '6', '₇': '7', '₈': '8', '₉': '9',

    # Superscripts
    'ᵃ': 'a', 'ᵇ': 'b', 'ᶜ': 'c', 'ᵈ': 'd', 'ᵉ': 'e',
    'ᶠ': 'f', 'ᵍ': 'g', 'ʰ': 'h', 'ᶦ': 'i', 'ʲ': 'j',
    'ᵏ': 'k', 'ˡ': 'l', 'ᵐ': 'm', 'ⁿ': 'n', 'ᵒ': 'o',
    'ᵖ': 'p', 'ʳ': 'r', 'ˢ': 's', 'ᵗ': 't', 'ᵘ': 'u',
    'ᵛ': 'v', 'ʷ': 'w', 'ˣ': 'x', 'ʸ': 'y', 'ᶻ': 'z',
}

GREEK_MAP = {
    # Uppercase and lowercase Greek letters normalized to lowercase
    'Α': 'alpha', 'α': 'alpha',
    'Β': 'beta', 'β': 'beta',
    'Γ': 'gamma', 'γ': 'gamma',
    'Δ': 'delta', 'δ': 'delta',
    'Ε': 'epsilon', 'ε': 'epsilon',
    'Ζ': 'zeta', 'ζ': 'zeta',
    'Η': 'eta', 'η': 'eta',
    'Θ': 'theta', 'θ': 'theta',
    'Ι': 'iota', 'ι': 'iota',
    'Κ': 'kappa', 'κ': 'kappa',
    'Λ': 'lambda', 'λ': 'lambda',
    'Μ': 'mu', 'μ': 'mu',
    'Ν': 'nu', 'ν': 'nu',
    'Ξ': 'xi', 'ξ': 'xi',
    'Ο': 'omicron', 'ο': 'omicron',
    'Π': 'pi', 'π': 'pi',
    'Ρ': 'rho', 'ρ': 'rho',
    'Σ': 'sigma', 'σ': 'sigma', 'ς': 'sigma',
    'Τ': 'tau', 'τ': 'tau',
    'Υ': 'upsilon', 'υ': 'upsilon',
    'Φ': 'phi', 'φ': 'phi',
    'Χ': 'chi', 'χ': 'chi',
    'Ψ': 'psi', 'ψ': 'psi',
    'Ω': 'omega', 'ω': 'omega'
}


# ------------------------
# Helper Functions
# ------------------------

def convert_unicode_subscripts(expr):
    """
    Converts Unicode sub/superscripts to ASCII and joins consecutive ones
    into a single underscore-prefixed identifier. E.g., Tₛₑₚ ᵣ → T_sepr
    """
    expr = expr.translate(SUBSCRIPT_MAP).translate(SUPERSCRIPT_MAP)

    result = []
    buffer = ''
    i = 0
    while i < len(expr):
        c = expr[i]
        if c in UNICODE_SUBSCRIPT_MAP:
            buffer += UNICODE_SUBSCRIPT_MAP[c]
        elif c == ' ' and i + 1 < len(expr) and expr[i + 1] in UNICODE_SUBSCRIPT_MAP:
            # skip space between subscripts/superscripts
            pass
        else:
            if buffer:
                result.append('_' + buffer)
                buffer = ''
            result.append(c)
        i += 1

    if buffer:
        result.append('_' + buffer)

    return ''.join(result)


def is_descriptive_block(expr):
    """
    Determines if an expression is a descriptive text block rather than a mathematical condition.
    Returns True if the expression contains multiple alphabetic words and lacks mathematical operators.
    """
    expr = expr.strip()
    # If it has more than 2 alphabetic words AND contains no math operators, treat as descriptive
    word_count = len(re.findall(r'\b[a-zA-Z]{2,}\b', expr))
    mathy = re.search(r'[<>=!+\-*/^()]', expr)
    return word_count >= 3 and not mathy


def is_descriptive_block_2(expr):
    """Alternative implementation to detect descriptive blocks"""
    expr = expr.strip()
    word_count = len(re.findall(r'\b[a-zA-Z]{2,}\b', expr))
    return word_count >= 3


def normalize_brackets(expr):
    """Standardizes different bracket types to parentheses"""
    return expr.replace('{', '(').replace('}', ')').replace('[', '(').replace(']', ')')


# ------------------------
# Enhanced Bracket Validation
# ------------------------

class BracketValidationError(Exception):
    """Custom exception for bracket validation errors"""

    def __init__(self, message, event_id=None, expression=None):
        self.message = message
        self.event_id = event_id
        self.expression = expression
        super().__init__(self.message)


def validate_brackets_detailed(expr, event_id=None):
    """
    Enhanced bracket validation that provides detailed error information.
    Returns tuple: (is_valid, error_message, error_position)
    """
    stack = []
    brackets = {'(': ')', '[': ']', '{': '}'}
    open_brackets = set(brackets.keys())
    close_brackets = set(brackets.values())
    reverse_brackets = {v: k for k, v in brackets.items()}

    for i, char in enumerate(expr):
        if char in open_brackets:
            stack.append((char, i))
        elif char in close_brackets:
            if not stack:
                return False, f"Unexpected closing bracket '{char}' at position {i + 1}", i

            last_open, last_pos = stack.pop()
            expected_close = brackets[last_open]

            if char != expected_close:
                return False, (f"Mismatched brackets: '{last_open}' at position {last_pos + 1} "
                               f"closed with '{char}' at position {i + 1}"), i

    if stack:
        unclosed = stack[-1]
        return False, f"Unclosed bracket '{unclosed[0]}' at position {unclosed[1] + 1}", unclosed[1]

    return True, None, None


def is_balanced(expr):
    """
    Enhanced version that uses detailed validation.
    Maintains backward compatibility with existing code.
    """
    is_valid, _, _ = validate_brackets_detailed(expr)
    return is_valid


# ------------------------
# Expression Normalization
# ------------------------

def normalize_expr(expr):
    """
    Normalizes a logical expression by:
    - Converting unicode characters to standard ASCII
    - Standardizing logical operators
    - Normalizing comparison operators
    - Removing units from numerical expressions
    - Cleaning up whitespace
    """
    # Convert unicode subscripts/superscripts and Greek letters
    expr = convert_unicode_subscripts(expr)
    for greek, rep in GREEK_MAP.items():
        expr = expr.replace(greek, rep)

    # Standardize logical operators (with word boundary checks)
    expr = re.sub(r'\b(AND)\b', 'AND', expr)
    expr = re.sub(r'\b(OR)\b', 'OR', expr)

    # Convert alternative logical operators
    expr = expr.replace('&&', ' AND ').replace('||', ' OR ')

    # Remove phrases like "any 2 of the following:"
    expr = re.sub(r'any\s+\d+\s+of\s+the\s+following.*?:', '', expr, flags=re.IGNORECASE)

    # Normalize comparison operators
    expr = expr.replace("≥", ">=").replace("≤", "<=").replace("≠", "!=")

    # Fix assignment operators used as comparisons
    expr = re.sub(r'(?<![=!<>])=(?!=)', '==', expr)

    # Remove whitespace around comparison operators
    expr = re.sub(r'\s*(>=|<=|!=|==|>|<)\s*', r'\1', expr)

    # Only remove units in mathematical expressions, not in descriptive blocks
    if not is_descriptive_block(expr):
        # Remove common units with careful boundary checking
        expr = re.sub(
            r'(\d+(?:\.\d+)?)\s+(km|m|ms|s|g|m\/s2|m\/s²|V|°C|Hz|sec|msec|mV|km\/h|N|kg|A|Ω|ohm|rad|deg|min|h|mph|psi|bar|Pa|W|kW|rpm|dB|mol|cd|lux|lm)\b',
            r'\1', expr, flags=re.IGNORECASE)

        # Handle units directly attached to numbers (no space)
        expr = re.sub(r'(\d+(?:\.\d+)?)(ms|s|g|V|Hz|N|kg|A|W|h)\b',
                      r'\1', expr, flags=re.IGNORECASE)

    # Clean up extra spaces
    expr = re.sub(r'\s+', ' ', expr).strip()

    # Remove arrow/implies if present
    expr = re.sub(r'^.*?[=:>→\-]>\s*', '', expr).strip()

    # Normalize timing variables (e.g., TContactor ON ADC N → TContactor_ON_ADC_N)
    def normalize_var(match):
        var = match.group(0)
        var = convert_unicode_subscripts(var)
        var = re.sub(r'[\s\-]+', '_', var)
        var = re.sub(r'[^\w]', '', var)
        return var

    expr = re.sub(r'\bT[A-Za-z0-9₀₁₂₃₄₅₆₇₈₉ᵣᵤᵥₙₘₛₜₓᵃ-ᶻ\s\-]{1,40}\b', normalize_var, expr)

    return expr


# ------------------------
# Recursive Expression Parser
# ------------------------

def split_on_operator(expr, operator, level):
    """
    Splits an expression on a specific logical operator while respecting bracket nesting.
    Only splits at the top level (depth = 0).
    """
    depth = 0
    parts = []
    current = ""
    i = 0

    while i < len(expr):
        # Check for operator pattern with proper word boundaries
        match_found = False
        if depth == 0 and i + len(operator) <= len(expr):
            # Exact match required with word boundaries
            if expr[i:i+len(operator)] == operator:
                is_left_boundary = (i == 0 or not expr[i-1].isalnum() and expr[i-1] != '_')
                is_right_boundary = (i+len(operator) == len(expr) or
                                    not expr[i+len(operator)].isalnum() and expr[i+len(operator)] != '_')

                if is_left_boundary and is_right_boundary:
                    parts.append(current.strip())
                    current = ""
                    i += len(operator)
                    match_found = True

        if not match_found:
            # Track bracket depth
            if expr[i] in "([{":
                depth += 1
            elif expr[i] in ")]}":
                depth -= 1

            current += expr[i]
            i += 1

    if current:
        parts.append(current.strip())

    return parts if len(parts) > 1 else [expr]


def parse_expression(expr, level=1, event_id=None):
    """
    Recursively parses a logical expression into a list of conditions and operators.
    Handles nested expressions with proper precedence.
    Now includes bracket validation with event tracking.
    """
    if is_descriptive_block(expr):
        return [{'type': 'condition', 'value': expr}]

    expr = expr.strip()

    # Validate brackets before parsing
    is_valid, error_msg, error_pos = validate_brackets_detailed(expr, event_id)
    if not is_valid:
        raise BracketValidationError(
            f"Bracket validation failed: {error_msg}",
            event_id=event_id,
            expression=expr
        )

    # Handle parenthesized expression groups
    if (expr.startswith('(') and expr.endswith(')')) or (expr.startswith('[') and expr.endswith(']')):
        # Check if these are just enclosing brackets
        inner = expr[1:-1].strip()
        if is_balanced(inner):
            # Increase level for expressions inside brackets
            return parse_expression(inner, level + 1, event_id)

    # Try to split on top-level OR (lower precedence)
    parts = split_on_operator(expr, 'OR', level)
    if len(parts) > 1:
        result = []
        for i, part in enumerate(parts):
            result.extend(parse_expression(part, level, event_id))
            if i < len(parts) - 1:
                result.append({'type': 'operator', 'value': 'OR', 'level': level})
        return result

    # Try to split on top-level AND (higher precedence)
    parts = split_on_operator(expr, 'AND', level)
    if len(parts) > 1:
        result = []
        for i, part in enumerate(parts):
            result.extend(parse_expression(part, level, event_id))
            if i < len(parts) - 1:
                result.append({'type': 'operator', 'value': 'AND', 'level': level})
        return result

    # No top-level operators, so this is a basic condition
    return [{'type': 'condition', 'value': expr}]


def parse_logical_expression(expr, event_id=None):
    """
    Main parsing function that:
    1. Validates bracket balance first
    2. Normalizes the expression format
    3. Parses it into a structured representation
    4. Formats the result for CSV/Excel/JSON output
    """
    # First, validate brackets in the original expression
    is_valid, error_msg, error_pos = validate_brackets_detailed(expr, event_id)
    if not is_valid:
        raise BracketValidationError(
            f"Invalid bracket structure in expression: {error_msg}",
            event_id=event_id,
            expression=expr
        )

    expr = normalize_expr(expr)
    expr = normalize_brackets(expr)

    # Validate again after normalization
    is_valid, error_msg, error_pos = validate_brackets_detailed(expr, event_id)
    if not is_valid:
        raise BracketValidationError(
            f"Invalid bracket structure after normalization: {error_msg}",
            event_id=event_id,
            expression=expr
        )

    # Remove outermost brackets if they're balanced
    if (expr.startswith('(') and expr.endswith(')')) or (expr.startswith('[') and expr.endswith(']')):
        inner = expr[1:-1].strip()
        if is_balanced(inner):
            expr = inner

    # Parse the expression
    result = parse_expression(expr, event_id=event_id)

    # Format for CSV
    formatted_conditions = []
    formatted_operators = []

    # Process the parsed result
    for i, item in enumerate(result):
        if item['type'] == 'condition':
            formatted_conditions.append(item['value'])
        else:
            operator = item['value']
            # Format operators with asterisks based on nesting level
            prefix = '*' * item['level']
            if operator == 'AND':
                formatted_operators.append(f"{prefix}and")
            else:  # OR
                formatted_operators.append(f"{prefix}or")

    return formatted_conditions, formatted_operators


def extract_time_assignment(expr):
    """
    Extracts and normalizes the time variable from expressions like:
    "T₂ => [T >= (T1 + 0.600 s)]"
    Returns a C-compatible variable name using underscores.
    """
    assign_match = re.search(r'^(.*?)=>', expr.strip())
    if assign_match:
        raw_time_var = assign_match.group(1).strip()
        normalized_var = convert_unicode_subscripts(raw_time_var)

        # Additional normalization: Replace whitespace and dashes with underscores
        normalized_var = re.sub(r'[\s\-]+', '_', normalized_var)
        # Remove any remaining non-alphanumeric characters (except underscore)
        normalized_var = re.sub(r'[^\w]', '', normalized_var)
        return normalized_var

    return None


def parse_mission_expression(seq_id, expression):
    """
    Parses a mission expression and formats it for output with an event ID.
    Now includes bracket validation and error handling.
    """
    if isinstance(seq_id, str):
        seq_id = seq_id.strip()
        event_match = re.search(r'\bE\d+\b', seq_id)
        if event_match:
            seq_id = event_match.group(0)

    lines = [line.strip() for line in expression.strip().splitlines() if line.strip()]
    time_line_idx = -1
    flags_inside = []

    # Find the line with =>
    for idx, line in enumerate(lines):
        if '=>' in line:
            time_line_idx = idx
            break

    # Extract flag assignments before the =>
    for line in lines[:time_line_idx]:
        if re.match(r'^[A-Za-z_][A-Za-z0-9_]*\s*=\s*[^=]+$', line):
            flags_inside.append(line.strip())

    # Extract time assignment ONLY from the line that contains =>
    time_var = None
    if 0 <= time_line_idx < len(lines):
        time_var = extract_time_assignment(lines[time_line_idx])

    # Use full expression for logical parsing
    full_expr = " ".join(lines)

    try:
        conditions, operators = parse_logical_expression(full_expr, event_id=seq_id)
    except BracketValidationError as e:
        # Re-raise with more context
        raise BracketValidationError(
            f"Bracket error in {seq_id}: {e.message}",
            event_id=seq_id,
            expression=full_expr
        )

    # Build output row
    row = {'Event Number': seq_id}
    if time_var:
        row['time_assignment'] = time_var
    if flags_inside:
        row['flags_inside'] = flags_inside

    row['condition1'] = conditions[0] if conditions else ""
    for i in range(len(operators)):
        row[f'operator{i+1}'] = operators[i]
        if i+1 < len(conditions):
            row[f'condition{i+2}'] = conditions[i+1]

    return row


# ------------------------
# Document Parsing Functions
# ------------------------

def convert_pdf_to_docx(pdf_path):
    """Convert PDF to DOCX while preserving table structure"""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        docx_path = tmp.name
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
    return docx_path


def extract_from_docx(path):
    """
    Extracts mission logic expressions and event descriptions from a Word document.
    Now includes bracket validation and error collection.
    """
    doc = Document(path)
    rows = []
    bracket_errors = []

    for table in doc.tables:
        for r in table.rows[1:]:  # Skip header row
            cells = r.cells
            if len(cells) >= 3:
                event = cells[0].text.strip()
                logic = cells[1].text.strip()
                event_desc = " ".join(cells[2].text.split()).strip()  # Third column = Events

                try:
                    parsed = parse_mission_expression(event, logic)
                    parsed["event_description"] = event_desc
                    rows.append(parsed)
                except BracketValidationError as e:
                    bracket_errors.append({
                        'event': e.event_id or event,
                        'error': e.message,
                        'expression': e.expression or logic
                    })
                    print(f"❌ Bracket Error in {e.event_id or event}: {e.message}")

    if bracket_errors:
        # Save bracket errors to a separate file for review
        with open("bracket_validation_errors.json", "w", encoding="utf-8") as f:
            json.dump(bracket_errors, f, indent=4, ensure_ascii=False)

        error_summary = f"Found {len(bracket_errors)} bracket validation errors:\n"
        for error in bracket_errors:
            error_summary += f"  - {error['event']}: {error['error']}\n"

        raise BracketValidationError(
            f"Bracket validation failed for {len(bracket_errors)} events. "
            f"Details saved to 'bracket_validation_errors.json'.\n{error_summary}",
            event_id="MULTIPLE",
            expression="See bracket_validation_errors.json"
        )

    return rows


def extract_from_pdf(path):
    """First convert PDF to DOCX, then extract using DOCX parser"""
    print("Converting PDF to DOCX...")
    docx_path = convert_pdf_to_docx(path)
    try:
        return extract_from_docx(docx_path)
    finally:
        os.remove(docx_path)


# ------------------------
# DOP/DIP Parser Components (for explicit-only extraction)
# ------------------------

# Name-based alias mapping to RT#
rt_name_to_rt = {
    "MIU-A": "RT#3",
    "MIU-B": "RT#3",
    "CCSC": "RT#3",
    "IAU": "RT#4",
    "CPIF": "RT#5",
    "FNC": "RT#6"
}

# Mapping from RT# to subsystem and bus number
rt_to_subsystem = {
    'RT#1': 'ss1', 'RT#2': 'ss2', 'RT#3': 'ss3',
    'RT#4': 'ss4', 'RT#5': 'ss5', 'RT#6': 'ss6',
    'RT#01': 'ss1', 'RT#02': 'ss2', 'RT#03': 'ss3',
    'RT#04': 'ss4', 'RT#05': 'ss5', 'RT#06': 'ss6'
}
subsystem_to_bus = {
    'ss1': 0, 'ss2': 0, 'ss3': 0,
    'ss4': 1, 'ss5': 1, 'ss6': 1
}


def parse_dip_data(text):
    pattern = r'DIP\s+([\d\s,&]+)\s+of\s+(\w+)\s*\(\s*RT\s*#\s*(\d+)\s*\)'
    match = re.search(pattern, text, re.IGNORECASE)
    if match:
        raw_dip = match.group(1)
        subsystem = match.group(2)
        rtno = f'RT#{int(match.group(3))}'
        dip_list = [d.strip() for d in re.split(r'[,&]', raw_dip) if d.strip()]
        return dip_list, subsystem, rtno
    return None, None, None


def extract_flags(remark_text):
    """
    Extracts variable assignments (flags) from remarks like 'flag = 1' or 'NAV2Cmd = 0.5'.
    Avoids interpreting DIP/DOP as flags.
    """
    flags = []
    assignment_pattern = r'\b([A-Za-z_][A-Za-z0-9_]*)\s*=\s*([^,;\n]+)'
    matches = re.findall(assignment_pattern, remark_text)

    for var, expr in matches:
        if not re.match(r'(RT#?\d+|DIP|DOP)', var, re.IGNORECASE):
            flags.append(f"{var.strip()} = {expr.strip()}")
    return flags


def parse_remark(remark):
    remark = remark.strip()
    remark = re.sub(r'–|—|-', ' ', remark)  # replace all dash types with space
    flags = extract_flags(remark)
    issue_null = "yes" if "issue null command" in remark.lower() else "no"

    dop_n, dop_r = [], []
    status_n, status_r = "NIL", "NIL"

    # Handle all DOP formats: DOP# X,Y (N|R) (ON|OFF) — in any order
    pattern_1 = re.findall(r'DOP\s*#\s*([\w\d,\s]+?)\s*\(\s*(N|R)\s*\)\s*[\(\[\-]?\s*(ON|OFF)\s*[\)\]\-]?', remark,
                           re.IGNORECASE)
    pattern_2 = re.findall(r'DOP\s*#\s*([w\d,\s]+?)\s*[\(\[\-]?\s*(ON|OFF)\s*[\)\]\-]?\s*\(\s*(N|R)\s*\)', remark,
                           re.IGNORECASE)

    # Combined handling for both (group first or status first)
    combined_matches = pattern_1 + [(dop, grp, status) for dop, status, grp in pattern_2]

    for dop_list, group, status in combined_matches:
        dop_values = [d.strip() for d in dop_list.split(',') if d.strip()]
        if group.upper() == 'N':
            dop_n.extend(dop_values)
            status_n = status.upper()
        else:
            dop_r.extend(dop_values)
            status_r = status.upper()

    # Handle inline style like: DOP# 12,13(N) , 87,88(R) (ON)
    inline_match = re.findall(r'([\w\d,\s]+?)\s*\(\s*(N|R)\s*\)(?:\s*,)?', remark)
    inline_status_match = re.search(r'\b(ON|OFF)\b', remark, re.IGNORECASE)

    if inline_match and inline_status_match:
        inline_status = inline_status_match.group(1).upper()
        for dop_group, group_type in inline_match:
            dop_values = [d.strip() for d in dop_group.split(',') if d.strip()]
            if group_type.upper() == "N":
                dop_n.extend(dop_values)
                status_n = inline_status
            elif group_type.upper() == "R":
                dop_r.extend(dop_values)
                status_r = inline_status

    # Fallback for lone DOPs: "DOP# B10" or malformed
    fallback_dops = re.findall(r'DOP\s*#\s*([A-Z]?\d+)', remark, re.IGNORECASE)
    existing = set(dop_n + dop_r)
    for dop in fallback_dops:
        if dop not in existing:
            dop_n.append(dop)

    # --- Extract RT# from either explicit RT# or alias name ---
    rt_str = None

    # First try direct RT# mention
    # rt_match = re.search(r'RT\s*#\s*(\d+)', remark)
    rt_match = re.search(r'\(?RT\s*#\s*0*(\d+)\)?', remark, re.IGNORECASE)
    if rt_match:
        rt_str = f"RT#{rt_match.group(1)}"
    else:
        # Try alias name lookup
        for name, rt in rt_name_to_rt.items():
            if name in remark:
                rt_str = rt
                break

    if rt_str:
        subsys = rt_to_subsystem.get(rt_str, "unknown")
        bus = subsystem_to_bus.get(subsys, "unknown")
    else:
        subsys = "unknown"
        bus = "unknown"

    return {
        "busno": f"bus_{bus}" if bus != "unknown" else "unknown",
        "subsys": subsys,
        "dop_n": ", ".join(sorted(set(dop_n))) if dop_n else "NIL",
        "dop_r": ", ".join(sorted(set(dop_r))) if dop_r else "NIL",
        "status_n": status_n if dop_n else "NIL",
        "status_r": status_r if dop_r else "NIL",
        "issueNullCommand": issue_null,
        "flags": flags if flags else "NIL"
    }


def extract_dop_dip_from_docx(filepath):
    """Extracts event-wise DOP/DIP remarks from a DOCX file without logic-based assumptions."""
    doc = Document(filepath)
    data = []

    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "remarks" in headers and "event" in headers[0]:
            remarks_idx = headers.index("remarks")
            event_idx = 0  # Event assumed in first column

            for row in table.rows[1:]:
                raw_event = row.cells[event_idx].text.strip()
                remark = row.cells[remarks_idx].text.strip()

                match = re.match(r'(E\d+)', raw_event)
                if match and remark:
                    event_id = match.group(1)
                    parsed = parse_remark(remark)
                    if parsed:
                        parsed["event"] = event_id
                        data.append(parsed)
    return data


def extract_dop_dip_from_pdf(filepath):
    """First convert PDF to DOCX, then extract using DOCX parser"""
    print("Converting PDF to DOCX...")
    docx_path = convert_pdf_to_docx(filepath)
    try:
        return extract_dop_dip_from_docx(docx_path)
    finally:
        os.remove(docx_path)


def process_dop_dip_data_by_event(filepath):
    """
    Extracts and parses DOP/DIP data by event (E1, E2, etc.) from a DOCX or PDF.
    Outputs results to Excel and JSON.
    """
    if filepath.lower().endswith(".docx"):
        data = extract_dop_dip_from_docx(filepath)
    elif filepath.lower().endswith(".pdf"):
        data = extract_dop_dip_from_pdf(filepath)
    else:
        raise ValueError("Only .docx and .pdf files are supported.")

    # Save to Excel and JSON
    df = pd.DataFrame(data)
    df.to_excel("parsed_dop_dip_by_event.xlsx", index=False)

    with open("parsed_dop_dip_by_event.json", "w") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

    print("✅ Event-wise DOP/DIP data parsing complete.")
    return df, data


# ------------------------
# Post-Processing Functions
# ------------------------

def apply_function_placeholders(parsed_data):
    """
    Transforms descriptive text blocks into function call placeholders.
    - Descriptive → conditionXforEventY()
    - DIP → processDips1EventY(), processDips2EventY(), etc.
    """
    transformed_data = []

    for event in parsed_data:
        new_event = {"Event Number": event["Event Number"]}
        event_id = event["Event Number"].replace("E", "")

        if "time_assignment" in event:
            new_event["time_assignment"] = event["time_assignment"]
        if "event_description" in event:
            new_event["event_description"] = event["event_description"]
        if "flags_inside" in event:
            new_event["flags_inside"] = event["flags_inside"]

        i = 1
        dip_counter = 1  # Track DIP conditions

        while f"condition{i}" in event:
            cond_key = f"condition{i}"
            operator_key = f"operator{i - 1}" if i > 1 else None
            original = event[cond_key]

            if "dip" in original.lower():
                new_event[cond_key] = f"processDips{dip_counter}Event{event_id}()  /* {original} */"
                dip_counter += 1
            elif is_descriptive_block_2(original):
                new_event[cond_key] = f"condition{i}forEvent{event_id}()  /* {original} */"
            else:
                new_event[cond_key] = original

            if operator_key and operator_key in event:
                new_event[operator_key] = event[operator_key]

            current_op_key = f"operator{i}"
            if current_op_key in event:
                new_event[current_op_key] = event[current_op_key]

            i += 1

        transformed_data.append(new_event)

    return transformed_data


def clean_parsed_data(data_frame):
    clean_json_data = []

    for _, row in data_frame.iterrows():
        clean_event = {"Event Number": row["Event Number"]}

        if 'flags_inside' in row and isinstance(row['flags_inside'], list) and row['flags_inside']:
            clean_event['flags_inside'] = row['flags_inside']

        if 'time_assignment' in row and pd.notna(row['time_assignment']):
            clean_event['time_assignment'] = row['time_assignment']

        if 'event_description' in row and pd.notna(row['event_description']):
            clean_event['event_description'] = row['event_description']

        condition_count = 1
        while f"condition{condition_count}" in row:
            condition_key = f"condition{condition_count}"
            operator_key = f"operator{condition_count - 1}" if condition_count > 1 else None

            if condition_key in row and pd.notna(row[condition_key]) and row[condition_key]:
                clean_event[condition_key] = row[condition_key]

                if operator_key and operator_key in row and pd.notna(row[operator_key]) and row[operator_key]:
                    clean_event[operator_key] = row[operator_key]

            current_operator_key = f"operator{condition_count}"
            if current_operator_key in row and pd.notna(row[current_operator_key]) and row[current_operator_key]:
                clean_event[current_operator_key] = row[current_operator_key]

            condition_count += 1

        clean_json_data.append(clean_event)

    return clean_json_data


# ------------------------
# Main Execution
# ------------------------

def process_mission_logic(filename):
    """Process a document for mission logic expressions"""
    print(f"\n=== Processing Mission Logic from {filename} ===")

    # Process the document based on file type
    if filename.lower().endswith(".docx"):
        parsed_rows = extract_from_docx(filename)
        print(f"Extracted mission logic data from DOCX file: {filename}")
    elif filename.lower().endswith(".pdf"):
        print("PDF detected - converting to DOCX first...")
        parsed_rows = extract_from_pdf(filename)
        print(f"Extracted mission logic data from PDF file: {filename}")
    else:
        raise ValueError(f"Unsupported file type: {filename}. Please upload a .docx or .pdf file.")

    # Create DataFrame from parsed rows
    df = pd.DataFrame(parsed_rows)

    # Clean data and create JSON output
    clean_json_data = clean_parsed_data(df)

    # Save clean JSON data
    with open("parsed_mission_logic_clean.json", "w", encoding="utf-8") as json_file:
        json.dump(clean_json_data, json_file, indent=4, ensure_ascii=False)
    print("✅ Clean JSON output saved to 'parsed_mission_logic_clean.json'")

    # Apply function placeholders transformation
    transformed_data = apply_function_placeholders(clean_json_data)

    # Save transformed JSON data
    with open("parsed_mission_logic_with_placeholders.json", "w", encoding="utf-8") as outfile:
        json.dump(transformed_data, outfile, indent=4, ensure_ascii=False)
    print("✅ Saved with function-call placeholders to 'parsed_mission_logic_with_placeholders.json'")

    # Save to CSV and Excel
    df.to_csv("parsed_mission_logic.csv", index=False)
    df.to_excel("parsed_mission_logic.xlsx", index=False)
    print("✅ Results saved to CSV and Excel files")

    # Display the DataFrame
    display(df)
    return df, clean_json_data


# Final combined structured JSON (merge logic and dop/dip by Event Number)
def save_final_structured_output(logic_data, dop_dip_data):
    # Create a map from event ID to DOP/DIP data
    dop_map = {
        item["event"]: {k: v for k, v in item.items() if k != "event"}
        for item in dop_dip_data
    }

    final_output = []

    for logic_entry in logic_data:
        event_no = logic_entry.get("Event Number")

        # Extract time assignment from the event data and use it
        time_assignment = logic_entry.get("time_assignment")

        # Build logic block (excluding Event Number and time_assignment)
        logic_block = {
            k: v for k, v in logic_entry.items()
            if k not in ("Event Number", "time_assignment", "event_description")
        }

        if "flags_inside" in logic_entry:
            logic_block["flags_inside"] = logic_entry["flags_inside"]

        # Build dop/dip block (updated to include issueNullCommand, remove dip)
        dop_block = dop_map.get(event_no, {
            "busno": "unknown",
            "subsys": "unknown",
            "dop_n": "NIL",
            "dop_r": "NIL",
            "status_n": "NIL",
            "status_r": "NIL",
            "issueNullCommand": "no",
            "flags": "NIL"
        })

        # Combine everything
        output_entry = {
            "Event Number": event_no,
            "event_description": logic_entry.get("event_description", ""),
            "time_assignment": time_assignment,
            "logic": logic_block,
            "dop_dip": dop_block
        }

        final_output.append(output_entry)

    # Write to JSON
    with open("final_mission_parsing_output.json", "w", encoding="utf-8") as f:
        json.dump(final_output, f, indent=4, ensure_ascii=False)

    print("✅ Final structured JSON saved to 'final_mission_parsing_output.json'")


def select_file():
    root = Tk()
    root.withdraw()  # Hide the main window
    file_path = filedialog.askopenfilename(
        title="Select a DOCX or PDF file",
        filetypes=[
            ("DOCX and PDF Files", "*.docx *.pdf"),
            ("Word Documents", "*.docx"),
            ("PDF Files", "*.pdf"),
        ]
    )
    if not file_path:
        print("❌ No file selected. Exiting.")
        exit()
    return file_path


if __name__ == "__main__":
    try:
        filename = select_file()

        if not filename:
            print("❌ No file selected.")
            exit()

        # Automatically extract both logic and DOP/DIP data
        mission_df, mission_data = process_mission_logic(filename)
        dop_dip_df, dop_dip_data = process_dop_dip_data_by_event(filename)

        # Final structured JSON merge
        if mission_data and dop_dip_data:
            transformed_data = apply_function_placeholders(mission_data)
            save_final_structured_output(transformed_data, dop_dip_data)

        print("\n✅ All processing complete!")

    except Exception as e:
        print(f"❌ Error: {e}")
